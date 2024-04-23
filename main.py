import tkinter as tk
from tkinter import messagebox
from paket.room import Room
from paket.apartment import Apartment
from paket.multistory_building import MultistoryBuilding
import docx
from openpyxl import Workbook

def calculate_and_save_report():
    building_type = building_type_var.get()

    if building_type == "Комната":
        length = float(length_entry.get())
        width = float(width_entry.get())
        height = float(height_entry.get())
        building = Room(length, width, height)
        total_area = building.calculate_area()
        heat_power = building.calculate_heat_power()
    elif building_type == "Квартира":
        length = float(length_entry.get())
        width = float(width_entry.get())
        height = float(height_entry.get())
        num_rooms = int(rooms_entry.get())
        building = Apartment(length, width, height, num_rooms)
        total_area = building.calculate_total_area()
        heat_power = building.calculate_heat_power()
    else:  # Многоэтажный дом
        length = float(length_entry.get())
        width = float(width_entry.get())
        height = float(height_entry.get())
        num_floors = int(floors_entry.get())
        num_units_per_floor = int(units_entry.get())
        building = MultistoryBuilding(length, width, height, num_floors, num_units_per_floor)
        total_area = building.calculate_total_area()
        heat_power = building.calculate_heat_power()

    # Отображение результатов
    result_label.config(text=f"Общая площадь: {total_area} кв.м\nТепловая мощность: {heat_power} Вт")

    # Сохранение результатов в отчет .docx
    doc = docx.Document()
    doc.add_heading('Результаты расчетов', level=1)
    doc.add_paragraph(f"Общая площадь: {total_area} кв.м")
    doc.add_paragraph(f"Тепловая мощность: {heat_power} Вт")
    doc.save('report.docx')
    
    # Сохранение результатов в отчет .xlsx
    wb = Workbook()
    ws = wb.active
    ws.append(["Общая площадь", "Тепловая мощность"])
    ws.append([total_area, heat_power])
    wb.save('report.xlsx')
    messagebox.showinfo("Сохранение", "Результаты сохранены в отчет.docx и отчет.xlsx")

# Создание GUI
root = tk.Tk()
root.title("Калькулятор строительства")

building_type_var = tk.StringVar()
building_type_var.set("Комната")

building_type_label = tk.Label(root, text="Выберите тип помещения:")
building_type_label.pack()

building_type_options = ["Комната", "Квартира", "Многоэтажный дом"]
building_type_menu = tk.OptionMenu(root, building_type_var, *building_type_options)
building_type_menu.pack()

length_label = tk.Label(root, text="Длина:")
length_label.pack()
length_entry = tk.Entry(root)
length_entry.pack()

width_label = tk.Label(root, text="Ширина:")
width_label.pack()
width_entry = tk.Entry(root)
width_entry.pack()

height_label = tk.Label(root, text="Высота:")
height_label.pack()
height_entry = tk.Entry(root)
height_entry.pack()

rooms_label = tk.Label(root, text="Количество комнат (если квартира):")
rooms_label.pack()
rooms_entry = tk.Entry(root)
rooms_entry.pack()

floors_label = tk.Label(root, text="Этажей (если многоэтажный дом):")
floors_label.pack()
floors_entry = tk.Entry(root)
floors_entry.pack()

units_label = tk.Label(root, text="Количество квартир на этаже (если многоэтажный дом):")
units_label.pack()
units_entry = tk.Entry(root)
units_entry.pack()

calculate_button = tk.Button(root, text="Рассчитать и сохранить", command=calculate_and_save_report)
calculate_button.pack()

result_label = tk.Label(root, text="")
result_label.pack()

root.mainloop()
